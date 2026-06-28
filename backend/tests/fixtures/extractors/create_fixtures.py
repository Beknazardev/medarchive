"""Create test fixtures for the extractors."""

import os
from pathlib import Path

FIXTURE_DIR = Path(__file__).parent


def create_html_fixtures():
    valid_html = b"""<!DOCTYPE html>
<html>
<head><title>Price List</title></head>
<body>
<h1>Medical Services</h1>
<p>Updated prices for 2026.</p>
<table id="prices">
<thead>
<tr><th>Service</th><th>Price (KZT)</th><th>Duration</th></tr>
</thead>
<tbody>
<tr><td>MRT head</td><td>15000</td><td>30 min</td></tr>
<tr><td>CT body</td><td>12000</td><td>20 min</td></tr>
<tr><td>Ultrasound</td><td>5000</td><td>15 min</td></tr>
</tbody>
</table>
<h2>Contact Information</h2>
<p>Phone: +7(712)1234567</p>
</body>
</html>"""
    (FIXTURE_DIR / "valid.html").write_bytes(valid_html)

    empty_table_html = b"""<!DOCTYPE html>
<html><body>
<table><tr><td></td><td></td></tr></table>
</body></html>"""
    (FIXTURE_DIR / "empty_table.html").write_bytes(empty_table_html)

    merged_headers_html = b"""<!DOCTYPE html>
<html><body>
<table>
<thead>
<tr><th colspan="2">Category A</th><th>Price</th></tr>
<tr><th>Service</th><th>Sub</th><th>Amount</th></tr>
</thead>
<tbody>
<tr><td>X-Ray</td><td>Basic</td><td>3000</td></tr>
</tbody>
</table>
</body></html>"""
    (FIXTURE_DIR / "merged_headers.html").write_bytes(merged_headers_html)

    malformed_html = b"not html at all <unclosed"
    (FIXTURE_DIR / "malformed.html").write_bytes(malformed_html)


def create_text_fixtures():
    valid_text = (
        "Service,Price,Currency\n"
        "MRT head,15000,KZT\n"
        "CT body,12000,KZT\n"
        "Ultrasound,5000,KZT\n"
    ).encode("utf-8")
    (FIXTURE_DIR / "valid.txt").write_bytes(valid_text)

    empty_text = b""
    (FIXTURE_DIR / "empty.txt").write_bytes(empty_text)

    single_line = b"just one line with data"
    (FIXTURE_DIR / "single_line.txt").write_bytes(single_line)


def create_xlsx_fixtures():
    import openpyxl
    from openpyxl.utils import get_column_letter

    wb = openpyxl.Workbook()
    ws = wb.active
    ws.title = "Prices"
    ws.append(["Service", "Price", "Currency", "Duration"])
    ws.append(["MRT head", 15000, "KZT", "30 min"])
    ws.append(["CT body", 12000, "KZT", "20 min"])
    ws.append(["Ultrasound", 5000, "KZT", "15 min"])

    ws2 = wb.create_sheet("Departments")
    ws2.append(["Department", "Floor"])
    ws2.append(["Radiology", "2"])
    ws2.append(["Cardiology", "3"])

    buf = BytesIO()
    wb.save(buf)
    (FIXTURE_DIR / "valid.xlsx").write_bytes(buf.getvalue())

    wb2 = openpyxl.Workbook()
    ws_empty = wb2.active
    ws_empty.title = "Empty"
    buf2 = BytesIO()
    wb2.save(buf2)
    (FIXTURE_DIR / "empty.xlsx").write_bytes(buf2.getvalue())


def create_docx_fixtures():
    from docx import Document

    doc = Document()
    doc.add_heading("Medical Price List", level=1)
    doc.add_paragraph("Updated for 2026.")

    table = doc.add_table(rows=4, cols=3)
    table.style = "Table Grid"
    table.cell(0, 0).text = "Service"
    table.cell(0, 1).text = "Price"
    table.cell(0, 2).text = "Duration"
    table.cell(1, 0).text = "MRT head"
    table.cell(1, 1).text = "15000"
    table.cell(1, 2).text = "30 min"
    table.cell(2, 0).text = "CT body"
    table.cell(2, 1).text = "12000"
    table.cell(2, 2).text = "20 min"
    table.cell(3, 0).text = "Ultrasound"
    table.cell(3, 1).text = "5000"
    table.cell(3, 2).text = "15 min"

    doc.add_paragraph("Contact: +7(712)1234567")

    buf = BytesIO()
    doc.save(buf)
    (FIXTURE_DIR / "valid.docx").write_bytes(buf.getvalue())

    doc2 = Document()
    doc2.add_paragraph("Empty document")
    buf2 = BytesIO()
    doc2.save(buf2)
    (FIXTURE_DIR / "minimal.docx").write_bytes(buf2.getvalue())


def create_pdf_fixtures():
    """Create a simple PDF using reportlab if available, else use a minimal PDF."""
    try:
        from reportlab.lib.pagesizes import A4
        from reportlab.pdfgen import canvas
        from reportlab.lib.units import cm

        buf = BytesIO()
        c = canvas.Canvas(buf, pagesize=A4)
        c.setFont("Helvetica", 12)
        c.drawString(2*cm, 28*cm, "Medical Price List")
        c.setFont("Helvetica", 10)
        c.drawString(2*cm, 26*cm, "Service: MRT head, Price: 15000 KZT")
        c.drawString(2*cm, 24*cm, "Service: CT body, Price: 12000 KZT")
        c.drawString(2*cm, 22*cm, "Service: Ultrasound, Price: 5000 KZT")
        c.save()
        (FIXTURE_DIR / "valid.pdf").write_bytes(buf.getvalue())
    except ImportError:
        # Create a minimal valid PDF manually
        minimal_pdf = _create_minimal_pdf()
        (FIXTURE_DIR / "valid.pdf").write_bytes(minimal_pdf)

    # Create a scanned-like PDF with minimal text
    scanned_pdf = _create_minimal_pdf(
        text="",
        title="Scanned Document",
    )
    (FIXTURE_DIR / "scanned.pdf").write_bytes(scanned_pdf)


def _create_minimal_pdf(text: str = "Service MRT head Price 15000\nService CT body Price 12000\n", title: str = "Price List") -> bytes:
    """Create a minimal valid PDF."""
    content = f"""%PDF-1.4
1 0 obj<</Type/Catalog/Pages 2 0 R>>endobj
2 0 obj<</Type/Pages/Kids[3 0 R]/Count 1>>endobj
3 0 obj<</Type/Page/Parent 2 0 R/MediaBox[0 0 612 792]/Contents 4 0 R/Resources<</Font<</F1 5 0 R>>>>>>endobj
4 0 obj
<</Length 144>>
stream
BT
/F1 12 Tf
50 750 Td
({title}) Tj
0 -20 Td
({text.replace(chr(10), ' ')}) Tj
ET
endstream
endobj
5 0 obj<</Type/Font/Subtype/Type1/BaseFont/Helvetica>>endobj
xref
0 6
0000000000 65535 f 
0000000009 00000 n 
0000000058 00000 n 
0000000115 00000 n 
0000000266 00000 n 
0000000462 00000 n 
trailer<</Size 6/Root 1 0 R>>
startxref
541
%%EOF"""
    return content.encode("latin-1")


from io import BytesIO


if __name__ == "__main__":
    print("Creating fixtures...")
    create_html_fixtures()
    create_text_fixtures()
    create_xlsx_fixtures()
    create_docx_fixtures()
    create_pdf_fixtures()
    print("Done. Created fixtures in", FIXTURE_DIR)
