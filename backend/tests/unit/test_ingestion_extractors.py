"""Tests for generic document extractors - Phase E RED tests."""

from __future__ import annotations

import pytest
from pathlib import Path
from typing import Any

from app.ingestion.extractors.base import (
    ExtractionLimits,
    ExtractionOutput,
    GenericCell,
    GenericTable,
    GenericTextBlock,
    MalformedDocument,
    ManualReviewRequired,
    MIMEMismatch,
    PasswordProtected,
    UnsupportedFormat,
)
from app.ingestion.extractors.html import extract_html
from app.ingestion.extractors.pdf import extract_pdf
from app.ingestion.extractors.docx import extract_docx
from app.ingestion.extractors.excel import extract_excel
from app.ingestion.extractors.text import extract_text

FIXTURE_DIR = Path(__file__).parent.parent / "fixtures" / "extractors"


# ─── Base contract tests ───

class TestExtractionLimits:
    def test_default_limits_are_reasonable(self):
        lim = ExtractionLimits()
        assert lim.max_pages == 50
        assert lim.max_sheets == 50
        assert lim.max_rows_per_table == 10_000
        assert lim.max_cols_per_table == 100
        assert lim.max_cell_length == 10_000
        assert lim.max_processing_seconds > 0

    def test_enforce_cell_length_truncates(self):
        lim = ExtractionLimits(max_cell_length=10)
        result = lim.enforce_cell_length("a" * 20)
        assert len(result) == 10

    def test_enforce_cell_length_keeps_short_text(self):
        lim = ExtractionLimits(max_cell_length=100)
        result = lim.enforce_cell_length("short")
        assert result == "short"


class TestGenericTable:
    def test_auto_counts_rows_and_cols(self):
        from app.ingestion.extractors.base import GenericRow
        cell1 = GenericCell(text="a", row_index=0, col_index=0)
        cell2 = GenericCell(text="b", row_index=0, col_index=1)
        cell3 = GenericCell(text="c", row_index=1, col_index=0)
        row1 = GenericRow(cells=(cell1, cell2), row_index=0)
        row2 = GenericRow(cells=(cell3,), row_index=1)
        table = GenericTable(rows=(row1, row2), table_index=0)
        assert table.row_count == 2
        assert table.col_count == 2


# ─── HTML Extractor Tests ───

class TestHTMLExtractor:
    def test_extracts_tables_from_valid_html(self):
        content = (FIXTURE_DIR / "valid.html").read_bytes()
        output = extract_html(content)
        assert len(output.tables) == 1
        table = output.tables[0]
        assert table.row_count >= 3  # header + 3 data rows
        assert table.col_count == 3

    def test_preserves_cell_text(self):
        content = (FIXTURE_DIR / "valid.html").read_bytes()
        output = extract_html(content)
        header_row = output.tables[0].rows[0]
        assert header_row.cells[0].text == "Service"
        assert header_row.cells[1].text == "Price (KZT)"
        data_row = output.tables[0].rows[1]
        assert data_row.cells[0].text == "MRT head"
        assert data_row.cells[1].text == "15000"

    def test_provenance_includes_source_url(self):
        content = (FIXTURE_DIR / "valid.html").read_bytes()
        output = extract_html(content, source_url="https://example.com/prices")
        assert output.tables[0].provenance.get("source_url") == "https://example.com/prices"

    def test_extracts_text_blocks(self):
        content = (FIXTURE_DIR / "valid.html").read_bytes()
        output = extract_html(content)
        assert len(output.text_blocks) >= 2  # h1 + p + h2 + p

    def test_handles_merged_headers(self):
        content = (FIXTURE_DIR / "merged_headers.html").read_bytes()
        output = extract_html(content)
        assert len(output.tables) == 1
        table = output.tables[0]
        assert table.row_count >= 3
        header_row = table.rows[0]
        merged_cells = [c for c in header_row.cells if c.is_merged]
        assert len(merged_cells) >= 1

    def test_handles_empty_table(self):
        content = (FIXTURE_DIR / "empty_table.html").read_bytes()
        output = extract_html(content)
        # Empty table cells should have empty text
        if output.tables:
            for row in output.tables[0].rows:
                for cell in row.cells:
                    assert cell.text == ""

    def test_respects_max_rows_limit(self):
        html = "<html><body><table>" + "<tr><td>x</td></tr>" * 100 + "</table></body></html>"
        output = extract_html(html.encode(), limits=ExtractionLimits(max_rows_per_table=5))
        assert len(output.tables[0].rows) <= 5

    def test_respects_max_cols_limit(self):
        html = "<html><body><table><tr>" + "<td>x</td>" * 200 + "</tr></table></body></html>"
        output = extract_html(html.encode(), limits=ExtractionLimits(max_cols_per_table=3))
        assert all(len(row.cells) <= 3 for row in output.tables[0].rows)

    def test_truncates_long_cell_text(self):
        long_text = "x" * 500
        html = f"<html><body><table><tr><td>{long_text}</td></tr></table></body></html>"
        output = extract_html(html.encode(), limits=ExtractionLimits(max_cell_length=100))
        assert output.tables[0].rows[0].cells[0].text == "x" * 100

    def test_rejects_huge_html(self):
        huge = b"<html><body>" + b"x" * 60_000_000 + b"</body></html>"
        with pytest.raises(MalformedDocument):
            extract_html(huge)

    def test_empty_output_for_no_tables(self):
        html = b"<html><body><p>No tables here</p></body></html>"
        output = extract_html(html)
        assert len(output.tables) == 0
        assert len(output.text_blocks) >= 1


# ─── PDF Extractor Tests ───

class TestPDFExtractor:
    def test_extracts_text_from_valid_pdf(self):
        content = (FIXTURE_DIR / "valid.pdf").read_bytes()
        output = extract_pdf(content)
        all_text = " ".join(b.text for b in output.text_blocks)
        assert "Price" in all_text or "15000" in all_text

    def test_scanned_pdf_returns_manual_review(self):
        content = (FIXTURE_DIR / "scanned.pdf").read_bytes()
        output = extract_pdf(content)
        assert output.manual_review_required is True
        assert any(e.code == "MANUAL_REVIEW_REQUIRED" for e in output.errors)

    def test_provenance_includes_page(self):
        content = (FIXTURE_DIR / "valid.pdf").read_bytes()
        output = extract_pdf(content, source_url="https://example.com/doc.pdf")
        if output.text_blocks:
            assert "page_index" in output.text_blocks[0].provenance

    def test_respects_max_pages(self):
        content = (FIXTURE_DIR / "valid.pdf").read_bytes()
        output = extract_pdf(content, limits=ExtractionLimits(max_pages=1))
        assert len(output.errors) == 0 or any(
            "Page limit" in str(e) for e in output.errors
        )

    def test_rejects_too_small_content(self):
        with pytest.raises(MalformedDocument):
            extract_pdf(b"PK")

    def test_metadata_contains_counts(self):
        content = (FIXTURE_DIR / "valid.pdf").read_bytes()
        output = extract_pdf(content)
        assert "total_text_length" in output.metadata
        assert output.metadata["total_text_length"] >= 0


# ─── DOCX Extractor Tests ───

class TestDOCXExtractor:
    def test_extracts_table_from_valid_docx(self):
        content = (FIXTURE_DIR / "valid.docx").read_bytes()
        output = extract_docx(content)
        assert len(output.tables) >= 1
        table = output.tables[0]
        assert table.row_count >= 4  # header + 3 data rows

    def test_preserves_cell_text(self):
        content = (FIXTURE_DIR / "valid.docx").read_bytes()
        output = extract_docx(content)
        header_row = output.tables[0].rows[0]
        assert header_row.cells[0].text == "Service"
        data_row = output.tables[0].rows[1]
        assert data_row.cells[0].text == "MRT head"
        assert data_row.cells[1].text == "15000"

    def test_extracts_paragraphs(self):
        content = (FIXTURE_DIR / "valid.docx").read_bytes()
        output = extract_docx(content)
        assert len(output.text_blocks) >= 1
        all_text = " ".join(b.text for b in output.text_blocks)
        assert "2026" in all_text

    def test_provenance_includes_source_url(self):
        content = (FIXTURE_DIR / "valid.docx").read_bytes()
        output = extract_docx(content, source_url="https://example.com/doc.docx")
        assert output.tables[0].provenance.get("source_url") == "https://example.com/doc.docx"

    def test_minimal_docx_no_tables(self):
        content = (FIXTURE_DIR / "minimal.docx").read_bytes()
        output = extract_docx(content)
        assert len(output.tables) == 0
        assert len(output.text_blocks) >= 1

    def test_rejects_too_small_content(self):
        with pytest.raises(MalformedDocument):
            extract_docx(b"PK")

    def test_respects_max_rows_limit(self):
        from docx import Document
        from io import BytesIO
        doc = Document()
        table = doc.add_table(rows=20, cols=3)
        for i, row in enumerate(table.rows):
            for cell in row.cells:
                cell.text = f"cell_{i}"
        buf = BytesIO()
        doc.save(buf)
        output = extract_docx(buf.getvalue(), limits=ExtractionLimits(max_rows_per_table=5))
        assert len(output.tables[0].rows) <= 5


# ─── Excel Extractor Tests ───

class TestExcelExtractor:
    def test_extracts_all_sheets(self):
        content = (FIXTURE_DIR / "valid.xlsx").read_bytes()
        output = extract_excel(content)
        assert len(output.tables) >= 2  # Prices + Departments sheets

    def test_extracts_first_sheet_data(self):
        content = (FIXTURE_DIR / "valid.xlsx").read_bytes()
        output = extract_excel(content)
        prices_sheet = output.tables[0]
        assert prices_sheet.page_or_sheet == "Prices"
        assert prices_sheet.row_count >= 4  # header + 3 data rows

    def test_preserves_numeric_as_text(self):
        content = (FIXTURE_DIR / "valid.xlsx").read_bytes()
        output = extract_excel(content)
        data_row = output.tables[0].rows[1]
        assert data_row.cells[1].text == "15000"

    def test_provenance_includes_sheet_name(self):
        content = (FIXTURE_DIR / "valid.xlsx").read_bytes()
        output = extract_excel(content, source_url="https://example.com/data.xlsx")
        assert output.tables[0].provenance.get("source_url") == "https://example.com/data.xlsx"
        assert output.tables[0].page_or_sheet == "Prices"

    def test_empty_sheet(self):
        content = (FIXTURE_DIR / "empty.xlsx").read_bytes()
        output = extract_excel(content)
        assert len(output.tables) == 0

    def test_rejects_too_small_content(self):
        with pytest.raises(MalformedDocument):
            extract_excel(b"PK")

    def test_respects_max_sheets(self):
        import openpyxl
        from io import BytesIO
        wb = openpyxl.Workbook()
        for i in range(10):
            ws = wb.create_sheet(f"Sheet{i}")
            ws.append([f"data{i}"])
        buf = BytesIO()
        wb.save(buf)
        output = extract_excel(buf.getvalue(), limits=ExtractionLimits(max_sheets=3))
        assert len(output.tables) <= 3

    def test_truncates_long_cell_text(self):
        import openpyxl
        from io import BytesIO
        wb = openpyxl.Workbook()
        ws = wb.active
        ws.append(["x" * 500])
        buf = BytesIO()
        wb.save(buf)
        output = extract_excel(buf.getvalue(), limits=ExtractionLimits(max_cell_length=100))
        assert output.tables[0].rows[0].cells[0].text == "x" * 100


# ─── Text Extractor Tests ───

class TestTextExtractor:
    def test_extracts_lines_from_valid_text(self):
        content = (FIXTURE_DIR / "valid.txt").read_bytes()
        output = extract_text(content)
        assert len(output.text_blocks) == 4  # header + 3 data lines

    def test_first_line_is_header(self):
        content = (FIXTURE_DIR / "valid.txt").read_bytes()
        output = extract_text(content)
        assert output.text_blocks[0].text == "Service,Price,Currency"

    def test_preserves_cell_text(self):
        content = (FIXTURE_DIR / "valid.txt").read_bytes()
        output = extract_text(content)
        assert "MRT head" in output.text_blocks[1].text

    def test_provenance_includes_source_url(self):
        content = (FIXTURE_DIR / "valid.txt").read_bytes()
        output = extract_text(content, source_url="https://example.com/data.txt")
        assert output.text_blocks[0].provenance.get("source_url") == "https://example.com/data.txt"

    def test_empty_text(self):
        content = (FIXTURE_DIR / "empty.txt").read_bytes()
        output = extract_text(content)
        assert len(output.text_blocks) == 0

    def test_single_line(self):
        content = (FIXTURE_DIR / "single_line.txt").read_bytes()
        output = extract_text(content)
        assert len(output.text_blocks) == 1

    def test_rejects_non_utf8(self):
        # Latin-1 encoded bytes that aren't valid UTF-8
        with pytest.raises(MIMEMismatch):
            extract_text(b"\xff\xfe\x00\x01")

    def test_metadata_contains_line_count(self):
        content = (FIXTURE_DIR / "valid.txt").read_bytes()
        output = extract_text(content)
        assert output.metadata["line_count"] == 4


# ─── Cross-cutting: limits and determinism ───

class TestCrossCutting:
    def test_all_extractors_return_immutable_output(self):
        for extractor_fn, fixture_name in [
            (extract_html, "valid.html"),
            (extract_docx, "valid.docx"),
            (extract_excel, "valid.xlsx"),
            (extract_text, "valid.txt"),
        ]:
            content = (FIXTURE_DIR / fixture_name).read_bytes()
            output = extractor_fn(content)
            assert isinstance(output, ExtractionOutput)

    def test_all_extractors_handle_zero_limits_gracefully(self):
        for extractor_fn, fixture_name in [
            (extract_html, "valid.html"),
            (extract_docx, "valid.docx"),
            (extract_excel, "valid.xlsx"),
            (extract_text, "valid.txt"),
        ]:
            content = (FIXTURE_DIR / fixture_name).read_bytes()
            output = extractor_fn(content, limits=ExtractionLimits(max_rows_per_table=1, max_cols_per_table=1, max_sheets=1))
            assert isinstance(output, ExtractionOutput)

    def test_all_extractors_produce_deterministic_output(self):
        content = (FIXTURE_DIR / "valid.html").read_bytes()
        out1 = extract_html(content)
        out2 = extract_html(content)
        assert out1.tables == out2.tables
        assert out1.text_blocks == out2.text_blocks
