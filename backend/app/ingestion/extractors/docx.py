"""DOCX extractor - extracts tables and paragraphs from Word documents."""

from __future__ import annotations

import time
from io import BytesIO
from typing import Any

from docx import Document

from app.ingestion.extractors.base import (
    ExtractionLimits,
    ExtractionOutput,
    GenericCell,
    GenericRow,
    GenericTable,
    GenericTextBlock,
    MalformedDocument,
    check_timeout,
)


def extract_docx(
    content: bytes,
    *,
    source_url: str | None = None,
    limits: ExtractionLimits | None = None,
) -> ExtractionOutput:
    lim = limits or ExtractionLimits()
    start = time.monotonic()
    errors: list[Any] = []
    tables: list[GenericTable] = []
    text_blocks: list[GenericTextBlock] = []

    if len(content) < 100:
        raise MalformedDocument("DOCX content too small", stage="docx")

    try:
        doc = Document(BytesIO(content))
    except Exception as exc:
        raise MalformedDocument(f"Cannot open DOCX: {exc}", stage="docx") from exc

    for tbl_idx, table in enumerate(doc.tables):
        check_timeout(start, lim.max_processing_seconds)
        if tbl_idx >= lim.max_sheets:
            break
        rows = _extract_docx_table(table, tbl_idx, lim)
        if rows:
            tables.append(
                GenericTable(
                    rows=tuple(rows),
                    table_index=tbl_idx,
                    page_or_sheet=f"table_{tbl_idx + 1}",
                    provenance={"source_url": source_url} if source_url else {},
                )
            )

    block_index = 0
    for para in doc.paragraphs:
        check_timeout(start, lim.max_processing_seconds)
        text = para.text.strip()
        if not text or len(text) < 3:
            continue
        truncated = lim.enforce_cell_length(text)
        style_name = (para.style.name or "Normal").lower() if para.style else "normal"
        text_blocks.append(
            GenericTextBlock(
                text=truncated,
                block_index=block_index,
                page_or_sheet="docx",
                block_type=style_name,
                provenance={"source_url": source_url} if source_url else {},
            )
        )
        block_index += 1

    return ExtractionOutput(
        tables=tuple(tables),
        text_blocks=tuple(text_blocks),
        errors=tuple(errors),
        metadata={"table_count": len(tables), "text_block_count": len(text_blocks)},
    )


def _extract_docx_table(
    table: Any,
    tbl_idx: int,
    limits: ExtractionLimits,
) -> list[GenericRow]:
    rows: list[GenericRow] = []
    for row_idx, row in enumerate(table.rows):
        if row_idx >= limits.max_rows_per_table:
            break
        cells: list[GenericCell] = []
        for col_idx, cell in enumerate(row.cells):
            if col_idx >= limits.max_cols_per_table:
                break
            text = cell.text.strip()
            truncated = limits.enforce_cell_length(text)
            cells.append(
                GenericCell(
                    text=truncated,
                    row_index=row_idx,
                    col_index=col_idx,
                    is_header=(row_idx == 0),
                )
            )
        if cells:
            rows.append(
                GenericRow(
                    cells=tuple(cells),
                    row_index=row_idx,
                    is_header=(row_idx == 0),
                )
            )
    return rows
